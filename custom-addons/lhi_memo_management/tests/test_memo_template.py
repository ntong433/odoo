# -*- coding: utf-8 -*-
import io
from unittest.mock import patch, MagicMock
from odoo.tests.common import TransactionCase
from odoo.exceptions import UserError, ValidationError
from ..services.word_template_service import WordTemplateService, REQUIRED_PLACEHOLDERS
from docx import Document


class TestMemoWordTemplate(TransactionCase):

    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        cls.company = cls.env.company
        cls.user = cls.env.user
        cls.department = cls.env["lhi.department"].create({
            "name": "Memo Template Test Department",
            "code": "MEMO-TEMPLATE-TEST",
            "company_id": cls.company.id,
        })

        cls.template = cls.env["lhi.memo.document.template"].create({
            "name": "LHI Internal Memo Template",
            "code": "LHI-INTERNAL-MEMO-TEST",
            "version": "1.0",
            "sharepoint_version": "1.0",
            "sharepoint_drive_id": "b!test_drive_id_12345",
            "sharepoint_item_id": "01TESTITEMID12345",
            "sharepoint_site_id": "test-site-id-12345",
            "sharepoint_web_url": "https://lhisokoto.sharepoint.com/sites/ERP/test.docx",
            "is_default": True,
            "active": True,
            "company_id": cls.company.id,
        })

        cls.category = cls.env["lhi.memo.category"].create({
            "name": "Test Category",
            "code": "TEST_CAT",
            "company_id": cls.company.id,
        })
        cls.connection = cls.env["lhi.graph.connection"].create({
            "name": "Memo Template Test Graph",
            "company_id": cls.company.id,
            "tenant_id": "11111111-1111-4111-8111-111111111111",
            "client_id": "22222222-2222-4222-8222-222222222222",
        })

    def _create_valid_template_bytes(self):
        doc = Document()
        doc.add_paragraph("REF: {{ memo_reference }}")
        doc.add_paragraph("FROM: {{ from_display }}")
        doc.add_paragraph("TO: {{ to_display }}")
        doc.add_paragraph("DATE: {{ memo_date }}")
        doc.add_paragraph("SUBJECT: {{ subject }}")
        doc.add_paragraph("{{ memo_body }}")
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()

    def test_01_single_active_default_template_constraint(self):
        """Only one active default template per company is allowed."""
        with self.assertRaises(ValidationError):
            self.env["lhi.memo.document.template"].create({
                "name": "Second Default Template",
                "code": "LHI-INTERNAL-MEMO-TEST-2",
                "version": "1.0",
                "sharepoint_version": "1.0",
                "sharepoint_drive_id": "b!test_drive_id_67890",
                "sharepoint_item_id": "01TESTITEMID67890",
                "is_default": True,
                "active": True,
                "company_id": self.company.id,
            })

    def test_02_template_snapshot_on_memo_creation(self):
        """Creating a memo assigns default template and snapshots metadata without auto-creating docx."""
        memo = self.env["lhi.memo"].create({
            "title": "Equipment Procurement Request",
            "subject": "Equipment Procurement Request",
            "purpose": "Request approval to procure equipment.",
            "memo_category_id": self.category.id,
            "requester_id": self.user.id,
            "department_id": self.department.id,
            "company_id": self.company.id,
        })

        self.assertEqual(memo.document_template_id, self.template)
        self.assertEqual(memo.template_version_snapshot, "1.0")
        self.assertEqual(memo.template_sharepoint_drive_id_snapshot, "b!test_drive_id_12345")
        self.assertEqual(memo.template_sharepoint_item_id_snapshot, "01TESTITEMID12345")
        self.assertEqual(memo.document_state, "not_created")
        self.assertFalse(memo.source_docx_item_id)

    def test_03_placeholder_validation_missing_placeholders(self):
        """WordTemplateService.validate_template raises UserError listing missing required placeholders."""
        doc = Document()
        output = io.BytesIO()
        doc.save(output)
        empty_bytes = output.getvalue()

        with self.assertRaises(UserError) as cm:
            WordTemplateService.validate_template(empty_bytes)

        error_msg = str(cm.exception)
        self.assertIn("missing the following required placeholders", error_msg)
        for ph in REQUIRED_PLACEHOLDERS:
            self.assertIn(ph, error_msg)

    def test_04_placeholder_validation_valid_template(self):
        """WordTemplateService.validate_template succeeds using get_undeclared_template_variables when all required placeholders exist."""
        valid_bytes = self._create_valid_template_bytes()
        self.assertTrue(WordTemplateService.validate_template(valid_bytes))

    def test_05_safe_filename_generation(self):
        """_safe_memo_filename converts slashes and sanitizes invalid SharePoint characters."""
        memo = self.env["lhi.memo"].create({
            "name": "LHI/MEMO/2026/00008",
            "title": "Network Equipment Request",
            "subject": 'Request for "Network" & : <Equipment>? / \\ |',
            "purpose": "Request approval for network equipment.",
            "memo_category_id": self.category.id,
            "requester_id": self.user.id,
            "department_id": self.department.id,
        })

        safe_name = memo._safe_memo_filename()
        self.assertTrue(safe_name.startswith("LHI-MEMO-2026-00008 - "))
        self.assertTrue(safe_name.endswith(".docx"))
        for invalid_char in ['"', "*", ":", "<", ">", "?", "/", "\\", "|"]:
            self.assertNotIn(invalid_char, safe_name)

    def test_06_validate_before_opening_word_missing_fields(self):
        """_validate_before_opening_word raises UserError listing missing required fields."""
        memo = self.env["lhi.memo"].create({
            "name": "New",
            "title": "Incomplete Memo",
            "subject": "Incomplete Memo",
            "purpose": "Exercise document preflight validation.",
            "memo_category_id": self.category.id,
            "requester_id": self.user.id,
            "department_id": self.department.id,
        })

        with self.assertRaises(UserError) as cm:
            memo._validate_before_opening_word()

        self.assertIn("required details are missing", str(cm.exception))

    def test_07_rendering_context_formatting(self):
        """_build_template_rendering_context formats dates, requester, recipients, and reference cleanly."""
        memo = self.env["lhi.memo"].create({
            "name": "LHI/MEMO/2026/00099",
            "title": "Test Context Memo",
            "subject": "Test Context Subject",
            "purpose": "<p>Test <b>body</b> content</p>",
            "memo_category_id": self.category.id,
            "requester_id": self.user.id,
            "department_id": self.department.id,
            "recipient_description": self.user.name,
        })

        context = memo._build_template_rendering_context()
        self.assertEqual(context["memo_reference"], "LHI/MEMO/2026/00099")
        self.assertEqual(context["subject"], "Test Context Subject")
        self.assertEqual(context["to_display"], self.user.name)
        self.assertIn("Test", context["memo_body"])

    def test_08_download_master_template_bytes_uses_binary_request_and_content_endpoint(self):
        """_download_master_template_bytes calls /content endpoint via lhi_binary_request with correct parameters."""
        memo = self.env["lhi.memo"].create({
            "title": "Download Test Memo",
            "subject": "Download Test Memo",
            "purpose": "Test the bounded template download path.",
            "memo_category_id": self.category.id,
            "requester_id": self.user.id,
            "department_id": self.department.id,
        })

        valid_docx = self._create_valid_template_bytes()
        mock_response = MagicMock()
        mock_response.content = valid_docx

        with patch.object(
            self.env.registry["lhi.graph.connection"],
            "lhi_binary_request",
            return_value=mock_response,
        ) as mock_binary:
            content = memo._download_master_template_bytes()
            self.assertTrue(content.startswith(b"PK"))
            mock_binary.assert_called_once()
            args, kwargs = mock_binary.call_args
            self.assertEqual(args[0], "GET")
            self.assertIn("/drives/", args[1])
            self.assertTrue(args[1].endswith("/content"))
            self.assertEqual(kwargs.get("auth_context"), "application")
            self.assertEqual(kwargs.get("expected_statuses"), {200})
            self.assertEqual(kwargs.get("allow_redirects"), True)

    def test_09_download_master_template_bytes_empty_content_raises_user_error(self):
        """Empty response content from SharePoint raises UserError."""
        memo = self.env["lhi.memo"].create({
            "title": "Empty Download Test Memo",
            "subject": "Empty Download Test Memo",
            "purpose": "Test rejection of an empty template download.",
            "memo_category_id": self.category.id,
            "requester_id": self.user.id,
            "department_id": self.department.id,
        })

        mock_response = MagicMock()
        mock_response.content = b""

        with patch.object(
            self.env.registry["lhi.graph.connection"],
            "lhi_binary_request",
            return_value=mock_response,
        ):
            with self.assertRaises(UserError) as cm:
                memo._download_master_template_bytes()
            self.assertIn("downloaded from SharePoint is empty", str(cm.exception))

    def test_10_ordinary_requester_can_open_word_without_direct_document_item_acl(self):
        """Ordinary memo requester without lhi.document.item ACL can open Word document safely."""
        employee_group = self.env.ref("lhi_security.group_lhi_employee")
        ordinary_user = self.env["res.users"].create({
            "name": "Ordinary Memo User",
            "login": "ordinary_memo_user",
            "email": "ordinary_memo_user@example.test",
            "group_ids": [(6, 0, [self.env.ref("base.group_user").id, employee_group.id])],
            "lhi_department_ids": [(6, 0, [self.department.id])],
        })

        memo = self.env["lhi.memo"].with_user(ordinary_user).create({
            "title": "Ordinary User Operations Request",
            "subject": "Ordinary User Operations Request",
            "purpose": "Test least-privilege Word document creation.",
            "memo_category_id": self.category.id,
            "requester_id": ordinary_user.id,
            "department_id": self.department.id,
            "recipient_user_ids": [(6, 0, [ordinary_user.id])],
            "recipient_description": ordinary_user.name,
        })

        valid_docx = self._create_valid_template_bytes()
        mock_response = MagicMock()
        mock_response.content = valid_docx

        def confirm_upload(documents):
            for document in documents:
                document.sudo().write({
                    "sharepoint_site_id": "template-test-site",
                    "sharepoint_drive_id": "template-test-drive",
                    "sharepoint_item_id": f"template-test-item-{document.id}",
                    "sharepoint_web_url": (
                        f"https://tenant.sharepoint.com/template-test-item-{document.id}"
                    ),
                    "storage_state": "available",
                    "upload_state": "completed",
                })
            return True

        with (
            patch.object(
                self.env.registry["lhi.graph.connection"],
                "lhi_binary_request",
                return_value=mock_response,
            ),
            patch.object(
                self.env.registry["lhi.document.item"],
                "action_upload",
                confirm_upload,
            ),
        ):
            action = memo.with_user(ordinary_user).action_open_word()
            self.assertEqual(action.get("type"), "ir.actions.act_url")
            self.assertTrue(memo.has_word_document)
            self.assertEqual(memo.document_state, "created")
